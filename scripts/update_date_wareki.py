"""
このスクリプトは、指定されたWord文書 (.docx) ファイル内の和暦の日付パターンを検索し、
今日の日付の和暦表記に置換します。

主な機能:
- コマンドライン引数として単一のWordファイルまたはディレクトリのパスを受け取ります。
- ディレクトリが指定された場合、`--pattern` オプションを使用して処理対象の.docxファイルをフィルタリングできます。
  (例: `python scripts/update_date_wareki.py /path/to/folder --pattern "*.docx"`)
- Word文書内の段落および表内のテキストから和暦の日付パターンを検出して置換します。
- 置換後の日付は、`japanera` ライブラリを使用して自動的に「令和XX年XX月XX日」形式（「元年」対応）で生成されます。
- 処理後、元のファイルを上書き保存します。
- `japanera` ライブラリが必要です。未インストールの場合はエラーメッセージが表示されます。
"""
import docx
import re
import os
from datetime import date, datetime
import argparse
import glob

# japaneraライブラリをインポート
try:
    from japanera import EraDate
except ImportError:
    print("エラー: 'japanera' ライブラリが見つかりません。")
    print("お手数ですが、ターミナルで以下のコマンドを実行してインストールしてください。")
    print("pip install japanera")
    exit()

def replace_date_pattern_in_doc(doc, new_date_str):
    """
    Word文書内で和暦の日付パターンに一致する箇所を、新しい日付文字列で置換します。
    置換が行われたかどうかを示す真偽値を返します。
    """
    # 年月日のバリエーションに対応するため、正規表現を少し柔軟にします。
    wareki_pattern = re.compile(r'(令和|平成|昭和|大正|明治)\s*(\d+|元)\s*年\s*\d+\s*月\s*\d+\s*日')
    found_and_replaced = False

    # 段落内のテキストを置換
    for p in doc.paragraphs:
        # run（書式設定ごとのテキスト断片）を結合して完全なテキストを取得
        full_text = "".join(run.text for run in p.runs)
        if wareki_pattern.search(full_text):
            new_text = wareki_pattern.sub(new_date_str, full_text)
            # runをクリアして新しいテキストを先頭に追加（書式は失われます）
            for i in range(len(p.runs)):
                p.runs[i].text = ""
            p.runs[0].text = new_text
            found_and_replaced = True

    # 表内のテキストを置換
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full_text = "".join(run.text for run in p.runs)
                    if wareki_pattern.search(full_text):
                        new_text = wareki_pattern.sub(new_date_str, full_text)
                        for i in range(len(p.runs)):
                            p.runs[i].text = ""
                        p.runs[0].text = new_text
                        found_and_replaced = True

    return found_and_replaced

def process_file(target_file_path, today_wareki_str):
    """
    単一のWordファイルを処理します。
    """
    try:
        print(f"--- ファイル '{os.path.basename(target_file_path)}' を処理中 ---")
        doc = docx.Document(target_file_path)
        
        was_replaced = replace_date_pattern_in_doc(doc, today_wareki_str)

        if was_replaced:
            doc.save(target_file_path)
            print(f"日付を '{today_wareki_str}' に置換し、上書き保存しました。")
        else:
            print("ファイル内に置換対象の日付パターンが見つかりませんでした。")

    except FileNotFoundError:
        print(f"エラー: ファイル '{target_file_path}' が見つかりません。")
    except Exception as e:
        print(f"エラーが発生しました: {e}")
    finally:
        print("-" * (len(os.path.basename(target_file_path)) + 14))


def main():
    """
    メイン処理。コマンドライン引数を解釈してファイルやディレクトリを処理します。
    """
    parser = argparse.ArgumentParser(
        description='Word文書 (.docx) 内の和暦の日付を今日の日付（または指定日）に更新します。',
        formatter_class=argparse.RawTextHelpFormatter
    )
    parser.add_argument('path', help='処理対象のWordファイルまたはディレクトリのパス。')
    parser.add_argument('--pattern', default='*.docx', help='ディレクトリ処理時のファイル名パターン。例: "送付状_*.docx" デフォルトは全ての.docxファイル ("*.docx")。')
    parser.add_argument('--date', help='指定する場合の日付 (形式: YYYY-MM-DD)。省略時は今日の日付になります。', default=None)
    args = parser.parse_args()

    target_path = args.path

    # 日付の決定
    if args.date:
        try:
            target_date = datetime.strptime(args.date, '%Y-%m-%d').date()
        except ValueError:
            print("エラー: 日付の形式が正しくありません。YYYY-MM-DD 形式で指定してください。")
            return
    else:
        target_date = date.today()

    # EraDateオブジェクトは標準のdateオブジェクトから生成
    era_date = EraDate.from_date(target_date)
    # era属性から元号名(漢字)を取得
    era_name = era_date.era.kanji
    # グレゴリオ暦の年から元号の開始年を引いて、元号の年を計算
    era_year = era_date.year - era_date.era.since.year + 1
    # 元号の年が1年の場合は「元年」とする
    era_year_str = "元" if era_year == 1 else str(era_year)
    today_wareki_str = f"{era_name}{era_year_str}年{era_date.month}月{era_date.day}日"
    
    print(f"設定する日付: {today_wareki_str}\n")

    files_to_process = []
    if os.path.isdir(target_path):
        # ディレクトリが指定された場合、指定されたパターンに一致するファイルをすべて取得
        search_pattern = os.path.join(target_path, args.pattern)
        files_to_process = glob.glob(search_pattern)
        # パターンに`.docx`が含まれていない場合も考慮し、docxファイルのみを対象とする
        files_to_process = [f for f in files_to_process if f.lower().endswith('.docx')]
        print(f"ディレクトリ '{target_path}' で、パターン '{args.pattern}' に一致する {len(files_to_process)} 個の .docx ファイルを処理します。")

    elif os.path.isfile(target_path) and target_path.lower().endswith('.docx'):
        # 単一の.docxファイルが指定された場合
        files_to_process.append(target_path)
    else:
        print(f"エラー: '{target_path}' は有効な .docx ファイルまたはディレクトリではありません。")
        return

    if not files_to_process:
        print("処理対象の .docx ファイルが見つかりませんでした。")
        return
    
    # 各ファイルを処理
    for file_path in files_to_process:
        # Wordが作成する一時ファイル (~$で始まるファイル) を無視する
        if not os.path.basename(file_path).startswith('~'):
            process_file(file_path, today_wareki_str)

    print("\nすべての処理が完了しました。")

if __name__ == '__main__':
    main()